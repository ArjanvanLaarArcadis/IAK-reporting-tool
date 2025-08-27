# -*- coding: utf-8 -*-
"""
Created on Fri Oct 11 13:54:24 2024

This script processes delivery lists, extracts relevant data from Excel files,
formats the data into Word documents, and embeds associated photos.

Original Author for v0: tersteer0528
Refactored: Sammie Knoppert (W. AGPT)
"""

import os
import glob
import pandas as pd
import docx
from docx.shared import Pt, RGBColor
from docx.enum.style import WD_STYLE_TYPE
from .utils import (
    load_config)
import logging
import time


def load_opleverlijst(filepath: str) -> pd.DataFrame:
    """
    Load the delivery list from an Excel file.

    Parameters:
        filepath (str): Path to the Excel file.

    Returns:
        pd.DataFrame: DataFrame containing the delivery list.
    """
    return pd.read_excel(filepath)


def extract_complexcode(object_code: str) -> str:
    """
    Generate the complex code from the object code.

    Parameters:
        object_code (str): The original object code.

    Returns:
        str: The complex code.
    """
    parts = object_code.split('-')
    complex_code = ''.join(parts[:3])  # Adjust slicing as necessary
    return complex_code


def load_ora(path_ora: str) -> pd.DataFrame:
    """
    Load ORA data from an Excel file.
 
    Parameters:
        path_ora (str): Path to the ORA Excel file.
 
    Returns:
        pd.DataFrame: Processed ORA DataFrame.
    """
    try:
        logging.info(f"Loading ORA data from: {path_ora}")
 
        # Get all sheet names and find the one that starts with "ORA"
        excel_file = pd.ExcelFile(path_ora)
        ora_sheet = next(
            (sheet for sheet in excel_file.sheet_names if sheet.startswith("ORA")), None
        )
 
        if ora_sheet is None:
            raise ValueError(f"No sheet starting with 'ORA' found in the Excel file [{path_ora}]")

        # The first 9 rows are skipped as they contain metadata. Further, the 11th row is dropped, it is a empty row below the header.
        ora = pd.read_excel(path_ora, sheet_name=ora_sheet, skiprows=list(range(9)) + [10], dtype=str)
        
        # Many cells (grayed) are left empty to indicate that the value is the same as the cell above. These are filled with the value from above.
        ora["Element"] = ora["Element"].ffill()
        ora["Bouwdeel"] = ora["Bouwdeel"].ffill()
        # Remark that the (blue) row of an "Element" is empty in all other columns, also the "Bouwdeel" column. Hence, the "Bouwdeel" column is 
        # filled with the value from above as well, which is incorrect. However, this has no influence on the final output, and is therefore not handled.

        # Additions are marked with a "(Ontbost)" or "(+)" or "+" in the "Bouwdeel" column.
        # These are removed to keep the "Bouwdeel" column clean.
        ora["Bouwdeel"] = ora["Bouwdeel"].str.replace(r"\(Ontbost\)|\(\+\)|\+", "", regex=True).str.strip()

        return ora
    except Exception as e:
        logging.error(f"Error loading ORA data: {e}")
        raise


def configure_document_styles(document: docx.Document, style_name: str, font_size: int) -> None:
    """
    Configure styles for the Word document.

    Parameters:
        document (docx.Document): The Word document object.
        style_name (str): Name of the new style.
        font_size (int): Font size for the style.
    """
    styles = document.styles
    cell_style = styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
    cell_style.font.underline = False
    cell_style.font.size = Pt(font_size)
    cell_style.font.name = 'Arial'
    cell_style.font.color.rgb = RGBColor(0, 0, 0)


def extract_relevant_ora_data(ora: pd.DataFrame) -> pd.DataFrame:
    """
    Extract relevant ORA data based on specific criteria.

    Parameters:
        ora (pd.DataFrame): The original ORA DataFrame.

    Returns:
        pd.DataFrame: Filtered ORA DataFrame.
    """
    idx_schades = ora[ora['Schade nummer'].notnull()].index
    idx_aandachtspunt = ora[ora['Advies mutatie I-ORA & Onderhoud'] == 'Aandachtspunt voor de volgende inspectie'].index
    idx_prestatie_do = ora[ora['Advies mutatie I-ORA & Onderhoud'] == 'Valt onder prestatiecontract / dagelijks onderhoud.'].index

    combined_indices = idx_schades.append([idx_aandachtspunt, idx_prestatie_do]).sort_values()
    return ora.loc[combined_indices]


def add_photos_to_document(
    document: docx.Document,
    measure: pd.Series,
    photo_loc: str,
    row_idx: int
) -> None:
    """
    Add photos to the Word document based on the measure data.

    Parameters:
        document (docx.Document): The Word document object.
        measure (pd.Series): A row from the ORA_MAATREGEL DataFrame.
        photo_loc (str): Directory location of photos.
        row_idx (int): Current row index in the Word table.
    """
    fotonummers = str(measure['Fotonummers'])
    fotonummers_list = fotonummers.split(',') if ',' in fotonummers else [fotonummers]

    table = document.tables[0]
    num_columns = len(table.columns)

    for i, foto_nummer in enumerate(fotonummers_list):
        foto_nummer = foto_nummer.strip()
        if foto_nummer and foto_nummer.lower() != 'nan':
            photo_path = glob.glob(os.path.join(photo_loc, f"**/{foto_nummer}.*"), recursive=True)
            if photo_path:
                # Calculate the target column index for the photo
                target_col = 6 + i  # Columns 6 and 7 for two photos
                if target_col < num_columns:
                    table.cell(row_idx, target_col).paragraphs[0].add_run().add_picture(photo_path[0], width=Pt(100))  # Adjust width as needed
                else:
                    print(f"Warning: Table does not have column index {target_col} for photo insertion.")


def create_word_document(template_path: str, objectnaam: str, objectcode: str) -> docx.Document:
    """
    Create and configure a Word document based on a template.

    Parameters:
        template_path (str): Path to the Word template.
        objectnaam (str): Name of the object.
        objectcode (str): Code of the object.

    Returns:
        docx.Document: Configured Word document.
    """
    document = docx.Document(template_path)
    
    # Configure styles
    configure_document_styles(document, 'Cell', 7)
    configure_document_styles(document, 'Cell2', 10)

    # Set header information
    if len(document.paragraphs) >= 2:
        document.paragraphs[0].text = f"Objectnaam: {objectnaam}"
        document.paragraphs[1].text = f"Topcode: {objectcode}"

    return document


def save_document(document: docx.Document, save_loc: str, objectcode: str) -> None:
    """
    Save the Word document to the specified location.

    Parameters:
        document (docx.Document): The Word document object.
        save_loc (str): Directory where the document will be saved.
        objectcode (str): Code of the object for naming the file.
    """
    save_path = os.path.join(save_loc, f"SSK-raming schades {objectcode}.docx")
    document.save(save_path)


def process_measure(
    document: docx.Document,
    measure: pd.Series,
    counter: int,
    idx: int,
    idx_schades: pd.Index,
    cell_style: str
) -> None:
    """
    Process each measure and add it to the Word document table.

    Parameters:
        document (docx.Document): The Word document object.
        measure (pd.Series): A row from the ORA_MAATREGEL DataFrame.
        counter (int): The row counter of table of the word document
        idx (int): Current idx of the row of the DataFrame.
        idx_schades (pd.Index): Indices of 'schades'.
        cell_style (str): Style name to apply to table cells.
    """
    table = document.tables[0]
    table.add_row()

    # Populate table cells
    table.cell(counter + 1, 0).text = str(measure['Element'].partition(',')[0])
    table.cell(counter + 1, 0).paragraphs[0].style = cell_style

    table.cell(counter + 1, 1).text = str(measure['Bouwdeel'])
    table.cell(counter + 1, 1).paragraphs[0].style = cell_style

    schade_omschrijving = (
        str(measure['Schade omschrijving']) if idx in idx_schades
        else str(measure['Bevinding:\n- Inspectie\n- Onderhoud\n- Overig'])
    )
    table.cell(counter + 1, 2).text = schade_omschrijving
    table.cell(counter + 1, 2).paragraphs[0].style = cell_style

    table.cell(counter + 1, 3).text = str(measure['Advies mutatie I-ORA & Onderhoud'])
    table.cell(counter + 1, 3).paragraphs[0].style = cell_style

    prestatie = str(measure['Maatregel omschrijving']) if idx in idx_schades else ""
    table.cell(counter + 1, 4).text = prestatie
    table.cell(counter + 1, 4).paragraphs[0].style = cell_style

    adviesjaar = (
        f"Adviesjaar: {str(measure['Optimaal JVA Onderhoud'])} \n"
        f"Uiterst jaar: {str(measure['Uiterst JVU Onderhoud'])}"
        if idx in idx_schades else ""
    )
    table.cell(counter + 1, 5).text = adviesjaar
    table.cell(counter + 1, 5).paragraphs[0].style = cell_style
