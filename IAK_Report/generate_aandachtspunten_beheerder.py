"""
This script generates "Bijlage 9 - Aandachtspunten Beheerder" documents based on data
from ORA (Onderhoudsrapportage) files. It processes relevant information, including
attention points and associated images, and formats it into a predefined Word template.
The final document is saved as both a Word file and a PDF.

The script performs the following steps:
1. Loads configuration and templates.
2. Iterates through objects, retrieving their paths and codes.
3. Extracts relevant data from the ORA file.
4. Processes the data to populate the Word document with attention points.
5. Saves the document and converts it to a PDF.

Functions:
- `create_word_document`: Creates a Word document based on a template and variables.
- `extract_relevant_data`: Filters the ORA data for relevant attention points.
- `list_of_fotonummers`: Parses and processes photo numbers from the data.
- `find_foto_path`: Finds the file path for a given photo number.
- `copy_last_table`: Duplicates the last table in the Word document.
- `remove_last_table`: Removes the last table in the Word document.
- `process_aandachtspunten_beheerder`: Populates the Word document with attention points.
- `save_aandachtspunten_beheerder`: Saves the Word document to a specified location.
- `main`: Orchestrates the entire process, including error handling.

Dependencies:
- `docx`: For creating and manipulating Word documents.
- `pandas`: For handling tabular data from the ORA file.
- `os`: For file and directory operations.
- `src.utils`: Custom utility functions for configuration, image handling, and PDF conversion.
- `src.get_voortgang`: Retrieves progress parameters for objects.
- `src.ora_to_word`: Loads ORA data into a DataFrame.

Usage:
Run this script directly to generate "Bijlage 9 - Aandachtspunten Beheerder" documents
for all objects in the batch listed in config.
"""
# Built-in modules
import os
import time
import logging
import copy
import datetime as dt

# External modules
import pandas as pd
import docx
from docx.shared import Pt, RGBColor
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT

# Local imports
from .utils import (
    load_config,
    get_object_paths_codes,
    convert_docx_to_pdf,
    list_pictures_for_object,
    update_config_with_voortgang,
    return_most_recent_ora,
    setup_logger,
    save_document,
)
from .get_voortgang import get_voortgang, get_voortgang_params
from .ora_to_word import load_ora


def create_word_document(template_path: str, variables: dict) -> docx.Document:
    """
    Create and configure a Word document based on a template provided by Rijkswaterstaat.

    Parameters:
        template_path (str): Path to the Word template.
        variables (dict): Dictionary containing configuration variables such as:
            - "complex_code" (str): Code of the complex.
            - "object_code" (str): Code of the object.
            - "object_naam" (str): Name of the object.

    Returns:
        docx.Document: Configured Word document.
    """
    logging.info("Creating Word document from template: %s", template_path)

    complex_code = variables.get("complex_code", "")
    object_code = variables.get("object_code", "")
    object_naam = variables.get("object_naam", "")

    logging.debug(
        "Loaded variables - complex_code: %s, object_code: %s, object_naam: %s",
        complex_code,
        object_code,
        object_naam,
    )

    document = docx.Document(template_path)
    styles = document.styles

    # Configure footer style
    if "FooterStyle" not in styles:
        logging.debug("Adding 'FooterStyle' to document styles.")
        footer_style = styles.add_style("FooterStyle", WD_STYLE_TYPE.PARAGRAPH)
        footer_style.font.underline = False
        footer_style.font.size = Pt(7)
        footer_style.font.name = "Arial"
        footer_style.font.color.rgb = RGBColor(0, 0, 0)
    else:
        logging.debug("'FooterStyle' already exists in document styles.")
        footer_style = styles["FooterStyle"]

    # Configure paragraph style
    if "Paragraph" not in styles:
        logging.debug("Adding 'Paragraph' style to document styles.")
        paragraph_style = styles.add_style("Paragraph", WD_STYLE_TYPE.PARAGRAPH)
        paragraph_style.font.underline = False
        paragraph_style.font.size = Pt(10)
        paragraph_style.font.name = "Arial"
        paragraph_style.font.color.rgb = RGBColor(0, 0, 0)
    else:
        logging.debug("'Paragraph' style already exists in document styles.")
        paragraph_style = styles["Paragraph"]

    # Configure footer content
    try:
        Footer = document.sections[0].footer
        Footer.tables[0].cell(0, 1).text = complex_code
        Footer.tables[0].cell(0, 1).paragraphs[0].style = footer_style
        Footer.tables[0].cell(1, 1).text = f"{object_code} {object_naam}"
        Footer.tables[0].cell(1, 1).paragraphs[0].style = footer_style
        logging.info("Footer content successfully configured.")
    except Exception as e:
        logging.error("Failed to configure footer content: %s", e)
        raise

    return document


def extract_relevant_data(ORA: pd.DataFrame) -> pd.DataFrame:
    """
    Filters the input DataFrame to extract rows where the column
    "Advies mutatie I-ORA & Onderhoud" contains both "aandachtspunt" and "beheerder" (case-insensitive).

    Args:
        ORA (pd.DataFrame): The input DataFrame containing the data to filter.

    Returns:
        pd.DataFrame: A filtered DataFrame containing only the rows
        where "Advies mutatie I-ORA & Onderhoud" contains both "aandachtspunt" and "beheerder".
    """
    logging.info(
        "Filtering ORA DataFrame for rows containing 'aandachtspunt' and 'beheerder' (case-insensitive)."
    )
    relevant_columns = [column for column in ORA.columns if column.startswith('Categorie')]
    select_column = relevant_columns[0] if relevant_columns else "Advies mutatie I-ORA & Onderhoud"
    return ORA[
        ORA[select_column].str.contains(
            "aandachtspunt", case=False, na=False
        )
        & ORA[select_column].str.contains(
            "beheerder", case=False, na=False
        )
    ]


def list_of_fotonummers(fotonummers: str) -> list:
    """
    Converts a cell value from a pandas Series containing photo numbers into a list of photo numbers.

    This function processes a string representation of photo numbers, which may be:
    - A comma-separated string of photo numbers.
    - A single photo number.
    - The string "nan" (interpreted as no photo numbers).

    Args:
        fotonummers (str): A cell value from a pandas Series, expected to be a string representation of photo numbers.

    Returns:
        list: A list of photo numbers as strings. Returns an empty list if the input is "nan".

    Debug Logging:
        Logs the input value and the resulting list for debugging purposes.
    """

    logging.debug("Processing fotonummers: %s", fotonummers)
    fotonummers = str(fotonummers)
    # Check if ',' exists in the string and split accordingly
    if "," in fotonummers:
        # Split the string into a list of substrings and strip each element
        fotonummers_list = [item.strip() for item in fotonummers.split(",")]
    elif ";" in fotonummers:
        # Split the string into a list of substrings and strip each element
        fotonummers_list = [item.strip() for item in fotonummers.split(";")]
    elif fotonummers.strip() == "nan":
        logging.debug("Input is 'nan', returning an empty list.")
        return []
    else:
        # Otherwise, just create a list with a single stripped element
        fotonummers_list = [fotonummers.strip()]

    logging.debug("Resulting list of fotonummers: %s", fotonummers_list)
    return fotonummers_list  # list of photo numbers as strings (filename)


def find_foto_path(fotonummer: str, imgs: list) -> str:
    """
    Finds the file path of an image based on a given photo number.

    This function searches through the list of files in the specified directory
    and returns the full path of the first image file that contains the given
    photo number in its name. Both the photo number and image filenames are
    transformed to lowercase and stripped of spaces for comparison.
    
    If multiple images match the photo number, the smallest file (compressed version)
    is returned.

    Args:
        fotonummer (str): The photo number to search for in the image filenames.
        imgs (list): The list of fullfilenames of all available images.

    Returns:
        str: The full file path of the matching image (smallest if multiple found),
             or raises FileNotFoundError if no match is found.
    """
    fotonummer = fotonummer.lower().replace(" ", "")
    
    # Case with extension:
    fotonummer, ext = os.path.splitext(fotonummer)
    if ext: 
        if not ext in ['.png', '.jpg', 'jpeg']:
            raise ValueError(f"[{ext}], that's a weird extension! Try to repair the ORA sheets")
    # Continue with just the name
    
    # Collect all matching images
    matching_images = []
    
    # The last part of the full filename contains the foto numbers
    for fullfilename in imgs:
        # Get only the name of the file, excluding the extension
        filename = os.path.basename(fullfilename)
        name, ext = os.path.splitext(filename)

        # Ok, sometimes, only the number is provided, without the camera-prefix.
        # e.g. 9252 instead of DSCN9252
        if name.lower().endswith(fotonummer):
            matching_images.append(fullfilename)
    
    # If no matches found, raise error
    if not matching_images:
        common_path = os.path.commonpath(imgs)
        raise FileNotFoundError(
            f"Image with photonummer [{fotonummer}] not found in [{common_path}]."
        )
    
    # If multiple matches, return the smallest file (compressed version)
    if len(matching_images) > 1:
        smallest_image = min(matching_images, key=os.path.getsize)
        logging.debug(f"Found {len(matching_images)} photos for {fotonummer}, using smallest: {os.path.basename(smallest_image)}")
        return smallest_image
    
    # Single match found
    logging.debug(f"Found photo {os.path.basename(matching_images[0])}")
    return matching_images[0]


def copy_last_table(word_document: docx.Document) -> None:
    """
    Duplicates the last table in the Word document.

    This function takes the last table in the provided Word template document,
    creates a deep copy of it, and appends the copied table to the document.

    Args:
        word_document (docx.Document): The Word document object where the table will be duplicated.

    Returns:
        None
    """
    #logging.debug("Copying the last table in the Word document.")
    template = word_document.tables[-1]
    tbl = template._tbl
    new_tbl = copy.deepcopy(tbl)
    paragraph = word_document.add_paragraph()
    paragraph._p.addnext(new_tbl)
    #logging.debug("Successfully copied and appended the last table.")


def remove_last_table(word_document: docx.Document) -> None:
    """
    Removes the last table in the Word document.

    This function identifies the last table in the provided Word document
    and removes it from the document.

    Args:
        word_document (docx.Document): The Word document object from which the last table will be removed.

    Returns:
        None
    """
    logging.debug("Removing the last table in the Word document.")
    if word_document.tables:
        tbl = word_document.tables[-1]._element
        tbl.getparent().remove(tbl)
        logging.debug("Successfully removed the last table.")
    else:
        logging.warning("No tables found in the Word document to remove.")


def process_aandachtspunten_beheerder(
    word_document: docx.Document, ora_filtered: pd.DataFrame, path_imgs: list
) -> docx.Document:
    """
    Processes the aandachtspunten beheerder and populates the Word document with relevant data.

    Args:
        word_document (docx.Document): The Word document to populate.
        ora_filtered (pd.DataFrame): Filtered ORA data containing aandachtspunten.
        path_imgs (str): Path to the directory containing images.

    Returns:
        docx.Document: The updated Word document.
    """
    logging.info("Starting to process aandachtspunten beheerder.")
    cell_style = word_document.styles["Paragraph"]

    # Duplicate tables based on the number of aandachtspunten
    logging.debug("Duplicating tables for aandachtspunten.")
    if len(ora_filtered) == 1:
        remove_last_table(word_document)
        logging.info("Removed the last table for a single aandachtspunt.")
    else:
        for _ in range(len(ora_filtered) - 2):
            copy_last_table(word_document)
        logging.info(f"Duplicated tables for {len(ora_filtered) - 1} aandachtspunten.")

    # Populate each table with data
    for i, (idx, row) in enumerate(ora_filtered.iterrows()):
        #logging.debug("Processing row %d: %s", idx, row.to_dict())

        cell_content = row["Bevinding:\n- Inspectie\n- Onderhoud\n- Overig"]
        if not ":" in cell_content:
            raise ValueError(
                f"Cell content does not contain ':': {cell_content}. "
                "Please check the ORA sheet for correct formatting."
            )

        # Everything in front of the colon is the attention point
        aandachtspunt = cell_content.partition(":")[0].strip()
        # However, sometimes it has an introduction-sentence, so we take only the two characters before
        # the colon, which is the attention point number.
        if len(aandachtspunt) > 2:
            aandachtspunt = aandachtspunt[-2:]

        # Everything after the colon is the observation
        bevinding_ora = cell_content.partition(":")[2].strip()
        
        # Extract photo numbers and their paths
        foto_column = [column for column, value in row.items() if "Foto" in column][0]
        fotos = list_of_fotonummers(row[foto_column])
        foto1 = fotos[0] if fotos else None
        foto2 = fotos[1] if len(fotos) > 1 else None
        path_foto1 = find_foto_path(foto1, path_imgs) if foto1 else None
        path_foto2 = find_foto_path(foto2, path_imgs) if foto2 else None

        logging.info(
            f"Aandachtspunt: {aandachtspunt}, Foto1: {foto1}, Foto2: {foto2}"
        )

        word_document.tables[i].cell(0, 0).text = str("Aandachtspunt " + aandachtspunt)
        word_document.tables[i].cell(0, 0).paragraphs[0].style = cell_style
        word_document.tables[i].cell(1, 1).text = str(row["Element"].partition(",")[0])
        word_document.tables[i].cell(1, 1).paragraphs[0].style = cell_style
        word_document.tables[i].cell(2, 1).text = str(row["Bouwdeel"])
        word_document.tables[i].cell(2, 1).paragraphs[0].style = cell_style
        word_document.tables[i].cell(4, 0).text = str(bevinding_ora)
        word_document.tables[i].cell(4, 0).paragraphs[0].style = cell_style
        word_document.tables[i].cell(4, 0).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        relevant_columns = [column for column, value in row.items() if column.startswith('Categorie')]
        select_column = relevant_columns[0] if relevant_columns else "Advies mutatie I-ORA & Onderhoud"
        word_document.tables[i].cell(6, 0).text = str(row[select_column])
        word_document.tables[i].cell(6, 0).paragraphs[0].style = cell_style
        word_document.tables[i].cell(6, 0).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP

        if foto1:
            logging.debug(f"Adding Foto1 to table {i}.")
            word_document.tables[i].cell(4, 2).paragraphs[0].add_run().add_picture(
                path_foto1, width=2350000
            )
        if foto2:
            logging.debug(f"Adding Foto2 to table {i}.")
            # TO DO: add an return between the two photos
            word_document.tables[i].cell(6, 2).paragraphs[0].add_run().add_picture(
                path_foto2, width=2350000
            )
        logging.info(f"Processed single aandachtspunt {i + 1}.")
    logging.info("Finished processing aandachtspunten beheerder.")
    return word_document


def save_aandachtspunten_beheerder(document: docx.Document, save_dir: str, object_code: str) -> str:
    """
    Save the Word document (Bijlage 9 - Aandachtspunten Beheerder) to the specified location.

    Args:
        document (docx.Document): The Word document object to save.
        save_dir (str): Directory path where the document should be saved.
        object_code (str): Code of the object used in the filename.

    Returns:
        str: The full path of the saved document.
    """
    # Extract the save location and construct the file name
    file_name = f"Bijlage 9 - Aandachtspunten Beheerder {object_code}.docx"

    # Delegate saving to the save_document function
    save_document(document, save_dir, file_name)

    # Construct and return the full save path for reference (optional)
    return os.path.join(save_dir, file_name)



def main():
    """
    Main function to orchestrate the processing of the PI report.
    """
    # Generate timestamped log filename
    timestamp = dt.datetime.now().strftime("%Y%m%d-%H%M%S")
    log_filename = f"generate_aandachtspunten_beheerder_{timestamp}.log"
    
    logger = setup_logger(log_filename)
    logger.info("Starting the generation process for aandachtspunten beheerder.")
    config_path = "./config.json"
    config = load_config(config_path=config_path)

    template_dir = "./templates"
    TEMPLATE_WORD = os.path.join(template_dir, "FORMAT_Bijlage9_AandachtspuntBeheerder.docx")
    TEMPLATE_WORD_GEEN = os.path.join(template_dir, "FORMAT_Bijlage9_GeenAandachtspuntBeheerder.docx")

    # Check if both template files exist
    if not os.path.exists(TEMPLATE_WORD):
        logger.error("Template file not found: %s", TEMPLATE_WORD)
        raise FileNotFoundError(f"Template file not found: {TEMPLATE_WORD}")
    if not os.path.exists(TEMPLATE_WORD_GEEN):
        logger.error("Template file not found: %s", TEMPLATE_WORD_GEEN)
        raise FileNotFoundError(f"Template file not found: {TEMPLATE_WORD_GEEN}")
    logger.info("Template files set successfully.")
    
    # Load the voortgang data
    if not config.get("voortgangs_sheet"):
        raise KeyError("Voortgangs sheet file not found in config.")
    excelfile = config["voortgangs_sheet"]
    df_voortgang = get_voortgang(excelfile=excelfile, abbrev=False)

    list_of_object_codes = get_object_paths_codes(config_file=config_path)
    failed_objects = []

    for object_path, object_code in list_of_object_codes:
        logger.info(f"Processing object path: {object_path}, object code: {object_code}")
        
        voortgang = get_voortgang_params(df_voortgang=df_voortgang, bh_code=object_code)
        variables = update_config_with_voortgang(config, voortgang)
        save_dir = os.path.join(object_path, config.get("output_folder", ""))
        try:
            path_ora = return_most_recent_ora(object_path)
            print("Checking for images...")
            path_imgs = list_pictures_for_object(object_path)
            ora = load_ora(path_ora)
            ora_filtered = extract_relevant_data(ora)
            logger.info(f"The number of aandachtspunten voor beheerder is: {len(ora_filtered)}")
            
            if len(ora_filtered) == 0:
                logger.info("Making the word document with no aandachtspunten...")
                word_document = create_word_document(TEMPLATE_WORD_GEEN, variables)
            else:
                logger.info("Making the word document with aandachtspunten...")
                word_document = create_word_document(TEMPLATE_WORD, variables)
                word_document = process_aandachtspunten_beheerder(
                    word_document, ora_filtered, path_imgs
                )
            
            document_path = save_aandachtspunten_beheerder(word_document, save_dir, object_code)
            logging.info(f"Word document saved successfully at: {document_path}")
            time.sleep(1)
            pdf_document_path = convert_docx_to_pdf(document_path)
            logging.info(f"PDF document for object code: {object_code} at [{pdf_document_path}]")
        except Exception as e:
            failed_objects.append(object_code)
            logging.error(f"Failed to generate for object code: {object_code}. Error: {e}")

    if failed_objects:
        logger.error(f"Failed to process the following objects: {failed_objects}")
    else:
        logger.info("All objects processed successfully.")


if __name__ == "__main__":
    main()
