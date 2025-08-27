"""
This script processes ORA (Object Risk Analysis) data to identify the highest risks
for one batch and generates a Word document summarizing these risks. The script also
generates an excel document to be used during the conversations. This document
is then used in conversations with the asset owners to discuss the identified risks.

The script performs the following steps:
1. Loads configuration settings from an Excel file.
2. Retrieves paths and codes for objects to process.
3. Identifies the most recent ORA file for each object.
4. Extracts relevant risk data from the ORA files based on specific criteria.
5. Combines the extracted data into a single DataFrame.
6. Creates a Word document using a predefined template.
7. Populates the Word document with the highest risks for each object.
8. Saves the generated Word document to a specified location.

Functions:
- `extract_relevant_ora_data`: Filters ORA data to extract rows with high risk scores.
- `create_word_document`: Creates and configures a Word document based on a template.
- `save_document`: Saves the Word document to a specified directory.
- `process_hoogste_risico`: Adds individual risk data to the Word document table.
- `main`: Orchestrates the entire process from data loading to document generation.

Dependencies:
- pandas
- python-docx
- Custom modules: `ora_to_word`, `utils`

Usage:
Run the script as a standalone program to generate the Word document summarizing the highest risks.
"""

import os
import pandas as pd
import docx
import datetime as dt
from .ora_to_word import (
    load_ora,
    configure_document_styles,
)
from .utils import (
    load_config,
    get_object_paths_codes,
    return_most_recent_ora,
    setup_logger,
    save_document,
    convert_docx_to_pdf,
)
import logging
import time
from openpyxl import Workbook
from openpyxl.utils.dataframe import dataframe_to_rows
from openpyxl.styles import Font, PatternFill, Alignment
from openpyxl.worksheet.table import Table, TableStyleInfo


def extract_relevant_ora_data(ora: pd.DataFrame) -> pd.DataFrame:
    """
    Extract relevant ORA data based on specific criteria.

    Parameters:
        ora (pd.DataFrame): The original ORA DataFrame.

    Returns:
        pd.DataFrame: Filtered ORA DataFrame.
    """

    ora['Actuele Risicoscore'] = ora['Actuele Risicoscore'].fillna(0).astype(int)
    idx_risicos = ora[
        ora["Actuele Risicoscore"] >= 6
    ].index  # Filter rows with score >= 6
    return ora.loc[idx_risicos]


def create_word_document(template_path: str, werkpakket: str) -> docx.Document:
    """
    Create and configure a Word document based on a template.

    Parameters:
        template_path (str): Path to the Word template.
        werkpakket (str): Name of the work package.

    Returns:
        docx.Document: Configured Word document.
    """
    logging.info("Creating Word document from template: %s", template_path)

    try:
        document = docx.Document(template_path)
    except Exception as e:
        logging.error("Failed to load Word template: %s", e)
        raise

    # Configure styles
    logging.info("Configuring document styles.")
    configure_document_styles(document, 'Cell', 7)
    configure_document_styles(document, 'Cell2', 10)

    # Set header information
    if len(document.paragraphs) >= 2:
        logging.info("Setting header information in the document.")
        document.paragraphs[0].text = f"Werkpakket: {werkpakket}"
        document.paragraphs[1].text = (
            f"Datum: {pd.Timestamp.now().strftime('%d-%m-%Y')}"
        )
    else:
        logging.warning(
            "Document template does not have enough paragraphs for headers."
        )

    logging.info("Word document created and configured successfully.")
    return document

def process_hoogste_risico(
    document: docx.Document,
    measure: pd.Series,
    counter: int,
    cell_style: str
) -> None:
    """
    Process each risico and add it to the Word document table.

    Parameters:
        document (docx.Document): The Word document object.
        measure (pd.Series): A row from the ORA_MAATREGEL DataFrame.
        counter (int): The row counter of table of the word document
        idx (int): Current idx of the row of the DataFrame.
        cell_style (str): Style name to apply to table cells.
    """
    table = document.tables[0]
    table.add_row()

    # Populate table cells

    table.cell(counter + 1, 0).text = str(measure["object_code"])
    table.cell(counter + 1, 0).paragraphs[0].style = cell_style

    table.cell(counter + 1, 1).text = str(measure['Element'].partition(',')[0])
    table.cell(counter + 1, 1).paragraphs[0].style = cell_style

    table.cell(counter + 1, 2).text = str(measure['Bouwdeel'])
    table.cell(counter + 1, 2).paragraphs[0].style = cell_style

    table.cell(counter + 1, 3).text = str(measure['Actuele Risicoscore'])
    table.cell(counter + 1, 3).paragraphs[0].style = cell_style

    table.cell(counter + 1, 4).text = str(measure['Actueel Risiconiveau'])
    table.cell(counter + 1, 4).paragraphs[0].style = cell_style

    table.cell(counter + 1, 5).text = str(measure["Bureaustudie:\n- Instandhoudingsrapportages\n- Toestandsinpecties\n- Overig"])
    table.cell(counter + 1, 5).paragraphs[0].style = cell_style

    table.cell(counter + 1, 6).text = str(measure["Toelichting.1"])
    table.cell(counter + 1, 6).paragraphs[0].style = cell_style


def save_dataframe_to_excel(
    df: pd.DataFrame, save_location: str, batch_name: str, col_mapping: dict
) -> None:
    """
    Save the DataFrame to an Excel file with formatting.

    Parameters:
        df (pd.DataFrame): The DataFrame to save.
        save_location (str): The directory where the file will be saved.
        batch_name (str): The batch name to use in the file name.
        col_mapping (dict): A dictionary mapping original column names to new names.
    """

    # Filter and rename columns in the DataFrame
    df = df[list(col_mapping.keys())].rename(columns=col_mapping)
    df["Element"] = df["Element"].str.split(",").str[0]

    # Define the file path
    excel_file_path = os.path.join(save_location, f"{batch_name} Hoogste Risicos.xlsx")

    # Create a new workbook and select the active worksheet
    wb = Workbook()
    ws = wb.active
    ws.title = "Hoogste Risicos"

    # Write the DataFrame to the worksheet
    for r_idx, row in enumerate(
        dataframe_to_rows(df, index=False, header=True), start=1
    ):
        for c_idx, value in enumerate(row, start=1):
            ws.cell(row=r_idx, column=c_idx, value=value)

    # Apply formatting to the header row
    header_fill = PatternFill(
        start_color="FFA500", end_color="FFA500", fill_type="solid"
    )
    header_font = Font(bold=True)
    for cell in ws[1]:
        cell.fill = header_fill
        cell.font = header_font

    # Create a table for the data
    table = Table(displayName="HoogsteRisicosTable", ref=ws.dimensions)
    style = TableStyleInfo(
        name="TableStyleMedium7",
        showFirstColumn=False,
        showLastColumn=False,
        showRowStripes=False,
        showColumnStripes=True,
    )
    table.tableStyleInfo = style
    ws.add_table(table)

    # Apply formatting to the columns
    column_widths = {
        "A": 12,
        "B": 35,
        "C": 35,
        "D": 12,
        "E": 20,
        "F": 80,
        "G": 80,
    }
    for col, width in column_widths.items():
        ws.column_dimensions[col].width = width

    # Apply text wrapping to all cells
    for row in ws.iter_rows():
        for cell in row:
            cell.alignment = Alignment(wrap_text=True, vertical="top")

    # Save the workbook
    wb.save(excel_file_path)
    logging.info(
        "Excel file with formatted table saved successfully at: %s", excel_file_path
    )


def main():
    """
    Main function to execute the script.
    """
    # Generate timestamped log filename
    timestamp = dt.datetime.now().strftime("%Y%m%d-%H%M%S")
    log_filename = f"generate_hoogste_risicos_{timestamp}.log"
    
    # Constants and configurations
    logger = setup_logger(log_filename, logging.INFO)
    logger.info("Starting the script.")

    config = load_config("data\config.json")
    logger.info("Configuration loaded successfully.")

    ORA_TEMPLATE_PATH = os.path.join(
        config["path_data_hoogste-risico"], "FORMAT_hoogste-risico.docx"
    )

    # Check if template file exists
    if not os.path.exists(ORA_TEMPLATE_PATH):
        logger.error("Template file not found: %s", ORA_TEMPLATE_PATH)
        raise FileNotFoundError(f"Template file not found: {ORA_TEMPLATE_PATH}")

    logger.info("Template file validated successfully.")

    SAVE_LOCATION = os.path.join(config["path_batch"], config["batch"])
    logger.info("Paths for template and save location configured.")

    # List all directories in BATCH_PATH
    object_paths_codes = get_object_paths_codes()
    logger.info("Object paths and codes retrieved successfully.")

    df_hoogste_risicos = pd.DataFrame()

    for path_object, object_code in object_paths_codes:
        logger.info("Processing object code: %s", object_code)
        path_ora = return_most_recent_ora(path_object)
        logger.info("Most recent ORA file located: %s", path_ora)

        # Load and process ORA data
        ora = load_ora(path_ora)
        logger.info("ORA data loaded successfully for object code: %s", object_code)

        ora_risico = extract_relevant_ora_data(ora)
        logger.info("Relevant ORA data extracted for object code: %s", object_code)

        ora_risico["object_code"] = object_code
        df_hoogste_risicos = pd.concat([df_hoogste_risicos, ora_risico], ignore_index=True)

    logger.info(
        "All object codes processed. Total risks identified: %d",
        len(df_hoogste_risicos),
    )

    # Create and configure Word document
    document = create_word_document(ORA_TEMPLATE_PATH, config["batch"])
    logger.info("Word document created and configured.")

    # Process each measure and add to document
    for counter, risico in df_hoogste_risicos.iterrows():
        logger.debug("Adding risico to document: %s", risico.to_dict())
        process_hoogste_risico(
            document=document,
            measure=risico,
            counter=counter,
            cell_style='Cell'
        )
    logger.info("All risks added to the Word document.")

    # Save the document
    save_document(document, SAVE_LOCATION, f"{config['batch']} Hoogste Risicos.docx")
    logger.info("Document saved successfully at: %s", SAVE_LOCATION)
    time.sleep(1)
    convert_docx_to_pdf(
        os.path.join(SAVE_LOCATION, f"{config['batch']} Hoogste Risicos.docx"),
        os.path.join(SAVE_LOCATION, f"{config['batch']} Hoogste Risicos.pdf"),
    )

    # Save the DataFrame to an Excel file
    col_mapping = {
        "object_code": "Code object",
        "Element": "Element",
        "Bouwdeel": "Bouwdeel",
        "Actuele Risicoscore": "Actuele Risicoscore",
        "Actueel Risiconiveau": "Actueel Risiconiveau",
        "Bureaustudie:\n- Instandhoudingsrapportages\n- Toestandsinpecties\n- Overig": "Bureaustudie",
        "Toelichting.1": "Toelichting",
    }
    save_dataframe_to_excel(
        df_hoogste_risicos, SAVE_LOCATION, config["batch"], col_mapping
    )
if __name__ == "__main__":
    main()
